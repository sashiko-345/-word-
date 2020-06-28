# 简单的打开word，输入数据，关闭word

def chouqu(document, n, n1, type, iseasy):
    if (type == 0):
        df = pd.read_excel("题库（单选+多选+判断）.xlsx", sheet_name=0, usecols=[2, 3, 4, 5, 6, 7, 8, 9])
    elif (type == 1):
        df = pd.read_excel("题库（单选+多选+判断）.xlsx", sheet_name=1, usecols=[2, 3, 4, 5, 6, 7, 8, 9])
    else:
        df = pd.read_excel("题库（单选+多选+判断）.xlsx", sheet_name=2, usecols=[2, 3, 9])
    df = df.values.tolist()
    df1 = []
    df2 = []
    df3 = []
    for i in range(len(df)):
        if df[i][0] == '低':
            df1.append(df[i])
        elif df[i][0] == '中':
            df2.append(df[i])
        else:
            df3.append(df[i])

    ans = []
    for i in range(n):
        if (iseasy == 1):
            a = random.choice(df1)
            df1.remove(a)
        elif (iseasy == 2):
            a = random.choice(df2)
            df2.remove(a)
        else:
            a = random.choice(df3)
            df3.remove(a)
        if type == 0 or type == 1:
            ans.append(a[7])
        else:
            ans.append(a[2])
        if (type == 0):
            document.add_paragraph("%s、%s\nA、%s\nB、%s\nC、%s\nD、%s\n" % (str(n1 + i + 1), a[1], a[2], a[3], a[4], a[5]))
        if (type == 1):
            document.add_paragraph(
                "%s、%s\nA、%s\nB、%s\nC、%s\nD、%s\nE、%s\n" % (str(n1 + i + 1), a[1], a[2], a[3], a[4], a[5], a[6]))
        if (type == 2):
            document.add_paragraph("%s、%s\n" % (str(n1 + i + 1), a[1]))

    return ans


def test(num, a1, a2, a3, b1, b2, b3, c1, c2, c3):
    document = Document()
    ans = []
    total = 0
    para = document.add_paragraph('一、单选题')

    ans.append(chouqu(document, a1, total, 0, 1))
    total += a1
    ans.append(chouqu(document, a2, total, 0, 2))
    total += a2
    ans.append(chouqu(document, a3, total, 0, 3))
    total += a3

    para = document.add_paragraph('二、多选题')
    ans.append(chouqu(document, b1, total, 1, 1))
    total += b1
    ans.append(chouqu(document, b2, total, 1, 2))
    total += b2
    ans.append(chouqu(document, b3, total, 1, 3))
    total += b3
    para = document.add_paragraph('三、判断题\n')
    ans.append(chouqu(document, c1, total, 2, 1))
    total += c1
    ans.append(chouqu(document, c2, total, 2, 2))
    total += c2
    ans.append(chouqu(document, c3, total, 2, 3))
    total += c3

    # 设置word字体大小
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    document.save('test%d.docx' % num)
    return ans

def write_answer(answer,n):
    for i in range(n):
        ans = open('answers-%d.txt'%i, 'w')
        ans.write(str(answer[i]))
        ans.close()

import pandas as pd
import random
from docx import Document
from docx.shared import Pt

number=int(input("请问要出几套卷子？"))
print("题库里有九种题，分别是单选题（简单题、中等题、难题）、多选题（简单题、中等题、难题）、判断题（简单题、中等题、难题）")
print("请输入你想要的各种题目数")
solo=input("单选题(输入格式：简单题数 中等题数 难题数）：")
a1,a2,a3 = map(lambda x:int(x),solo.split())
mul=input("多选题(输入格式：简单题数 中等题数 难题数）：")
b1,b2,b3 = map(lambda x:int(x),mul.split())
jud=input("判断题(输入格式：简单题数 中等题数 难题数）：")
c1,c2,c3 = map(lambda x:int(x),jud.split())

answer = []
for j in range(0,number):
        answer.append(test(j,a1,a2,a3,b1,b2,b3,c1,c2,c3))

write_answer(answer,number)
rep=[number,a1,a2,a3,b1,b2,b3,c1,c2,c3]
re = open('record.txt','w+')
re.write(str(rep))